from docx import Document

# Create a new Word document
doc = Document()

# Title
doc.add_heading("🕷️ Crawlability Report – Etsy.com", level=1)
doc.add_paragraph("Project: Intelligent Web Crawler & Analyzer")
doc.add_paragraph("Team Role: Member 1 – Crawlability Specialist")
doc.add_paragraph("Assigned Website: https://www.etsy.com")
doc.add_paragraph("\n")

# Section: Objective
doc.add_heading("🔍 Objective", level=2)
doc.add_paragraph("To analyze Etsy’s robots.txt file, determine crawl permissions, and summarize the rules that govern automated web crawlers.")

# Section: Robots.txt Overview
doc.add_heading("📄 Robots.txt Overview", level=2)
doc.add_paragraph("The robots.txt file for Etsy is available at:")
doc.add_paragraph("https://www.etsy.com/robots.txt")

# Section: Allowed and Disallowed Paths
doc.add_heading("✅ Allowed and ❌ Disallowed Paths", level=2)
doc.add_paragraph("Disallowed Paths (Partial List):")
disallowed = [
    "/c/", "/listing/", "/search/", "/your/", "/cart/",
    "/checkout/", "/people/", "/favourite/", "/register/",
    "/account/", "/email/", "/inbox/", "/int/", "/translations/"
]
for path in disallowed:
    doc.add_paragraph(f"Disallow: {path}", style='List Bullet')

doc.add_paragraph("\nThese disallowed rules mostly block:")
doc.add_paragraph("- User-specific sections (/your/, /account/, /cart/)")
doc.add_paragraph("- Listings and product pages (/listing/)")
doc.add_paragraph("- Search results (/search/)")
doc.add_paragraph("- Registration and messaging")

doc.add_heading("✅ Allowed Paths (Partial List)", level=2)
allowed_paths = [
    "/mailinglist/email/", "/uk/mailinglist/email/", "/au/mailinglist/email/", "/ca/mailinglist/email/",
    "/de-en/mailinglist/email/", "/dk-en/mailinglist/email/", "/fi-en/mailinglist/email/", "/hk-en/mailinglist/email/",
    "/api/v3/ajax/bespoke/public/neu/specs/"
    ]
for path in allowed_paths:
    doc.add_paragraph(f"Allow: {path}", style='List Bullet')

# Section: Crawl Delay
doc.add_heading("🔁 Crawl Delay", level=2)
doc.add_paragraph("Not specified. No Crawl-delay directive was found. Crawlers must respect server load responsibly.")

# Section: Sitemap
doc.add_heading("🔗 Sitemap Links", level=2)
sitemaps = [
    "https://www.etsy.com/sitemaps.xml",
    "https://www.etsy.com/blog/sitemap.xml.gz",
    "https://www.etsy.com/sitemaps/shops/vip-seller.xml",
    "https://www.etsy.com/fi-en/sitemaps/shops/vip-seller.xml",
    "https://www.etsy.com/au/sitemaps/shops/vip-seller.xml"
]
for sitemap in sitemaps:
    doc.add_paragraph(sitemap, style='List Bullet')

# Section: Code Snippet
doc.add_heading("🧠 Crawling Logic Code Snippet", level=2)
code = """
import urllib.robotparser

rp = urllib.robotparser.RobotFileParser()
rp.set_url("https://www.etsy.com/robots.txt")
rp.read()

test_url = "https://www.etsy.com/shop/some-shop"
if rp.can_fetch("*", test_url):
    print("✅ Allowed to crawl:", test_url)
else:
    print("❌ Disallowed to crawl:", test_url)
"""
doc.add_paragraph(code)

# Section: Summary Table
doc.add_heading("🧾 Summary", level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Light List Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Feature'
hdr_cells[1].text = 'Status'

summary = [
    ("Crawl Allowed?", "Partially (many paths blocked)"),
    ("Crawl Delay", "Not Specified"),
    ("Sitemap Provided", "✅ Yes"),
    ("Disallowed Focus", "Listings, Search, User Data"),
    ("Allowed Focus", "Homepage, Some Categories"),
]

for feature, status in summary:
    row_cells = table.add_row().cells
    row_cells[0].text = feature
    row_cells[1].text = status

# Save the document
doc.save("etsy_crawlability_report.docx")
print("✅ Word document generated: etsy_crawlability_report.docx")
